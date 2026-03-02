import asyncio
import json
import logging
import os
import uuid
from datetime import datetime as _dt
from typing import Any, Dict, List, Optional, Set

import httpx

from fastapi import APIRouter, Body, File, HTTPException, UploadFile

from app.models.schemas import (
    ChatRequest,
    ChatResponse,
    CodeGenRequest,
    CodeGenResponse,
    ConversationCreate,
    ConversationUpdate,
    DeepResearchRequest,
    DeepResearchResponse,
    ErrorEntryCreate,
    ImportMessagesRequest,
    IntentionRequest,
    Message,
    MSAgentTaskListItem,
    MSAgentTaskStatus,
    PDFToSkillsRequest,
    SkillCreate,
    SkillUpdate,
)
from app.services.model_router import model_router
from app.services.manager import manager_agent
from app.services.proxy_manager import proxy_manager
from app.services.memory_retriever import memory_retriever
from app.services.vector_sync import vector_sync
from app.services.executor import skill_executor
from app.services.skill_loader import skill_loader
from app.services.ms_agent_service import ms_agent_service
from app.services.tool_registry import tool_registry
from app.services.code_indexer import code_indexer  # Phase 11 代码语义索引
from sqlmodel import Session, select, delete
from app.core.db import engine
from app.models.database import Conversation, ConversationMessage, CodeSymbol, ErrorEntry, PDFDocument, RuntimeSetting, Task, SkillMetadata

router = APIRouter()
logger = logging.getLogger(__name__)

# ---- asyncio.create_task 引用持有集合（防止 GC 静默取消后台任务）----
_background_tasks: Set[asyncio.Task] = set()

def _track_task(t: asyncio.Task) -> asyncio.Task:
    """注册后台 Task，完成后自动从集合移除。"""
    _background_tasks.add(t)
    t.add_done_callback(_background_tasks.discard)
    return t


# ============================================================
# 对话 / 意图
# ============================================================

# 工具调用最大循环轮次（防止 LLM 无限循环调用工具）
_MAX_TOOL_ROUNDS = 30

# Agent System Prompt（Phase 10 C-1 + Phase 11 E-3）
# 注入条件：有可用工具 且 消息列表中尚无 system 消息
_AGENT_SYSTEM_PROMPT = """\
你是 Watery AI Agent，一个拥有自主工具调用能力的智能助手。

## 你的核心能力

你可以通过 Tool Calling 调用注册在系统中的技能（Skills）来执行真实操作，包括但不限于：
- 🔍 **联网搜索** (`web_search`) — 搜索实时互联网信息
- 📝 **错题库管理** (`error_ledger_crud`) — 创建/查询/删除错误经验记录
- 💬 **对话摘要** (`conversation_summary`) — 读取指定日期的对话历史并生成日报
- 📢 **飞书推送** (`feishu_webhook`) — 发送消息到飞书群
- 🛠️ **技能管理** (`skill_crud`) — 创建/更新/删除技能（你可以自我进化！）
- 🐍 **执行代码** (`run_python_snippet`) — 运行 Python 代码片段
- 📄 **PDF 处理** (`pdf_extract_text`, `pdf_to_skills`) — 提取和转化 PDF 知识
- 🔬 **代码搜索** (`code_search`) — 在项目代码中语义搜索函数/类/方法的精确位置（文件路径+行号）
- ...以及更多你可以自行创建的技能

## 工作准则

1. **先想后做**：收到复杂任务时，先在回复中简要列出你的执行计划（几步、每步用什么工具），
   然后再开始调用工具。
2. **逐步执行**：每步调用一个工具，评估结果后再决定下一步。
3. **自我进化**：如果你发现自己缺少某个能力，可以：
   a. 先用 `web_search` 搜索相关信息
   b. 用 `skill_crud` 将新学到的知识注册为技能
   c. 然后直接使用新创建的技能
4. **知识沉淀**：如果过程中发现了有价值的经验教训，用 `error_ledger_crud` 记录下来。
5. **代码定位优先**：当你需要查看或修改项目代码时，优先使用 `code_search` 定位目标函数，
   而非读取整个文件。这可以节省约 90% 的 Token 消耗。
6. **合理收束**：当任务完成后，给出清晰的最终回复，不要多余地调用工具。

## 注意事项

- 每次对话最多可调用 {max_rounds} 次工具，请合理规划。
- 工具列表会动态更新 —— 你新创建的技能在下一轮就可以使用。
- 如果某个工具调用失败，分析原因后可以换一种方式重试。
"""


@router.post("/chat", response_model=ChatResponse, summary="直连聊天接口（支持 Tool Calling + 会话持久化）")
async def chat_endpoint(request: ChatRequest):
    """
    通过 ModelRouter 动态路由到合适的模型提供商。

    **支持两种模式：**
    - **旧模式**：传入 ``messages``，后端不持久化（完全兼容现有行为）
    - **新模式**：传入 ``conversation_id`` + 当前 user ``messages``，后端自动加载历史并持久化结果

    **Tool Calling 流程**（当技能库已有注册技能时自动启用）：
    1. 获取全部注册技能（SQLite），转换为 OpenAI tool definitions
    2. 将 tools 列表随请求一起发给 LLM
    3. LLM 返回 ``finish_reason=tool_calls`` 时自动执行技能并将结果注入上下文
    4. LLM 返回 ``finish_reason=stop`` 时直接返回

    工具库为空时直接返回纯文本（兼容现有行为）。
    """
    try:
        conv_id = request.conversation_id

        # 获取工具定义（TTL 缓存，不频繁查库）
        tools = tool_registry.get_tool_definitions()

        # ---- 1. 加载消息 ----
        if conv_id:
            # 新模式：从 DB 加载完整历史消息（含 Tool Calling 上下文）
            with Session(engine) as session:
                conv = session.get(Conversation, conv_id)
                if not conv:
                    raise HTTPException(status_code=404, detail=f"Conversation '{conv_id}' not found.")
                db_msgs = session.exec(
                    select(ConversationMessage)
                    .where(ConversationMessage.conversation_id == conv_id)
                    .order_by(ConversationMessage.seq.asc())
                ).all()

            from app.models.schemas import ToolCall as ToolCallSchema
            messages: list = []
            for m in db_msgs:
                msg = Message(
                    role=m.role,
                    content=m.content,
                    tool_call_id=m.tool_call_id,
                )
                if m.tool_calls_json:
                    try:
                        msg.tool_calls = [
                            ToolCallSchema(**tc) for tc in json.loads(m.tool_calls_json)
                        ]
                    except Exception:
                        pass
                messages.append(msg)

            # 前端递入的新 user 消息（只取 user 角色）
            if request.messages:
                for m in request.messages:
                    if m.role == "user":
                        messages.append(m)
        else:
            # 旧模式：前端传完整 messages
            if not request.messages:
                raise HTTPException(status_code=400, detail="必须提供 messages 或 conversation_id。")
            messages = list(request.messages)

        # D-2: 按模型上下文窗口截断历史，防止超出 token 限制
        _CONTEXT_LIMITS = {
            "ark-code-latest": 28000,
            "gemini-2.0-flash": 100000,
            "gemini-1.5-flash": 100000,
            "gemini-2.5-pro-exp-03-25": 200000,
        }
        _DEFAULT_CTX = 28000
        max_ctx = _CONTEXT_LIMITS.get(request.model or "", _DEFAULT_CTX)
        messages = _truncate_messages(messages, max_tokens=max_ctx)

        # 记录 Tool Calling 循环开始前的消息数，循环期间新增的消息才需写入 DB
        messages_before_loop_len = len(messages)
        all_tool_results: list = []
        last_response: Optional[ChatResponse] = None

        # 读取运行时配置的轮次上限（可在 Settings 界面动态调整，默认 _MAX_TOOL_ROUNDS）
        _max_rounds = int(get_runtime_setting("max_tool_rounds") or _MAX_TOOL_ROUNDS)

        # Phase 10/11 — Agent System Prompt 注入
        # 条件：有可用工具 且 消息列表中尚无 system 消息（不覆盖用户手动设置的 system prompt）
        if tools and not any(m.role == "system" for m in messages):
            messages.insert(
                0,
                Message(
                    role="system",
                    content=_AGENT_SYSTEM_PROMPT.format(max_rounds=_max_rounds),
                ),
            )
            messages_before_loop_len = len(messages)  # system prompt 计入初始消息数

        # ---- 2. Tool Calling 循环 ----
        for round_num in range(_max_rounds):
            # 每轮刷新工具列表（支持对话中动态注册的新技能）
            tools = tool_registry.get_tool_definitions()
            last_response = await model_router.generate(
                messages=messages,
                model=request.model,
                temperature=request.temperature,
                max_tokens=request.max_tokens,
                tools=tools if tools else None,
            )

            # 无工具调用，直接返回
            if not last_response.tool_calls:
                last_response.tool_results = all_tool_results or None
                last_response.conversation_id = conv_id
                # 持久化到 DB（新模式）
                if conv_id:
                    _persist_chat_turn(
                        conv_id=conv_id,
                        new_messages=list(messages[messages_before_loop_len:]) + [
                            Message(role="assistant", content=last_response.content)
                        ],
                    )
                return last_response

            logger.info(
                f"chat_endpoint: round {round_num + 1}, "
                f"{len(last_response.tool_calls)} tool call(s) requested."
            )

            # 将 assistant 的 tool_calls 消息追加到对话
            messages.append(
                Message(
                    role="assistant",
                    content=last_response.content,
                    tool_calls=last_response.tool_calls,
                )
            )

            # 逐一执行工具调用
            for tc in last_response.tool_calls:
                skill = tool_registry.get_tool_by_name(tc.function.name)

                if skill:
                    try:
                        params = json.loads(tc.function.arguments)
                    except json.JSONDecodeError:
                        params = {}
                        logger.warning(
                            f"chat_endpoint: invalid JSON arguments for tool "
                            f"'{tc.function.name}': {tc.function.arguments!r}"
                        )

                    exec_result = await skill_executor.run(
                        language=skill.language,
                        entrypoint=skill.entrypoint,
                        params=params,
                    )
                    logger.info(
                        f"chat_endpoint: tool '{tc.function.name}' "
                        f"status={exec_result.get('status')}"
                    )
                else:
                    exec_result = {
                        "status": "error",
                        "message": f"Tool '{tc.function.name}' not found in skill registry.",
                    }
                    logger.warning(
                        f"chat_endpoint: unknown tool '{tc.function.name}' requested by LLM."
                    )

                all_tool_results.append({
                    "tool_call_id": tc.id,
                    "tool_name": tc.function.name,
                    "result": exec_result,
                })

                # 将工具执行结果以 role=tool 消息写回对话
                messages.append(
                    Message(
                        role="tool",
                        content=json.dumps(exec_result, ensure_ascii=False),
                        tool_call_id=tc.id,
                    )
                )

        # 超出最大轮次，返回最后一次响应（附带工具结果）
        logger.warning(
            f"chat_endpoint: reached max tool rounds ({_max_rounds}), returning last response."
        )
        if last_response:
            last_response.tool_results = all_tool_results or None
            last_response.conversation_id = conv_id
            if conv_id:
                _persist_chat_turn(
                    conv_id=conv_id,
                    new_messages=list(messages[messages_before_loop_len:]),
                )
            return last_response

        raise HTTPException(status_code=500, detail="Tool calling loop exhausted without a response.")

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Chat endpoint error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ---- Phase 7 辅助函数 ----

def _estimate_tokens(text: str) -> int:
    """快速估算文本 token 数。规则：UTF-8 字节数 / 3（适用于中英混合内容）。"""
    if not text:
        return 0
    return max(1, len(text.encode("utf-8")) // 3)


def _truncate_messages(messages: list, max_tokens: int = 28000) -> list:
    """
    智能截断消息列表，确保总 token 数不超过 max_tokens。

    策略（按优先级）：
    1. 永远保留 role=system 消息
    2. 永远保留最近 4 条消息（约最近 2 轮对话）
    3. 从最早的非 system 消息开始删除，直到 token 总数 < max_tokens
    4. 在截断点插入一条系统提示说明省略情况
    """
    def _count(msgs: list) -> int:
        total = 0
        for m in msgs:
            total += 4  # role + 格式开销
            total += _estimate_tokens(m.content or "")
            if getattr(m, "tool_calls", None):
                try:
                    total += _estimate_tokens(
                        json.dumps([tc.model_dump() for tc in m.tool_calls])
                    )
                except Exception:
                    pass
        return total

    if _count(messages) <= max_tokens:
        return messages

    system_msgs = [m for m in messages if m.role == "system"]
    other_msgs  = [m for m in messages if m.role != "system"]

    min_keep = min(4, len(other_msgs))
    recent     = other_msgs[-min_keep:]
    candidates = other_msgs[:-min_keep]

    removed = 0
    while candidates and _count(system_msgs + candidates + recent) > max_tokens:
        candidates.pop(0)
        removed += 1

    result = list(system_msgs)
    if removed > 0:
        result.append(Message(
            role="system",
            content=f"[注意: 更早的 {removed} 条消息已因上下文长度限制被省略。以下是最近的对话。]",
        ))
    result.extend(candidates)
    result.extend(recent)
    logger.info(f"_truncate_messages: 已截断 {removed} 条早期消息，剩余 {len(result)} 条")
    return result


def _persist_chat_turn(conv_id: str, new_messages: list) -> None:
    """
    将本次对话轮次中产生的所有新消息批量写入 ConversationMessage 表，
    并更新 Conversation 的 message_count、updated_at 和 title（自动命名）。

    Args:
        conv_id:      目标会话 ID
        new_messages: 本轮新增的 Message 对象列表（user + tool 中间消息 + final assistant）
    """
    if not new_messages:
        return
    try:
        from datetime import datetime as _dt
        with Session(engine) as session:
            conv = session.get(Conversation, conv_id)
            if not conv:
                logger.warning(f"_persist_chat_turn: conversation '{conv_id}' not found, skipping.")
                return

            current_seq = conv.message_count

            for msg in new_messages:
                tc_json = None
                if getattr(msg, "tool_calls", None):
                    try:
                        tc_json = json.dumps(
                            [tc.model_dump() for tc in msg.tool_calls],
                            ensure_ascii=False,
                        )
                    except Exception:
                        tc_json = None

                db_msg = ConversationMessage(
                    conversation_id=conv_id,
                    role=msg.role,
                    content=msg.content,
                    tool_calls_json=tc_json,
                    tool_call_id=getattr(msg, "tool_call_id", None),
                    token_count=_estimate_tokens(msg.content or ""),
                    seq=current_seq,
                )
                session.add(db_msg)
                current_seq += 1

            conv.message_count = current_seq
            conv.updated_at = _dt.utcnow()

            # 自动更新标题（仍为默认"新对话"且有 user 消息时）
            if conv.title == "新对话":
                first_user = next((m for m in new_messages if m.role == "user"), None)
                if first_user and first_user.content:
                    raw = first_user.content
                    conv.title = raw[:20] + ("..." if len(raw) > 20 else "")

            session.add(conv)
            session.commit()
    except Exception as e:
        logger.error(f"_persist_chat_turn failed for conv '{conv_id}': {e}")


# ============================================================
# 聊天附件上传（图片/PDF/文档 → 提取可供 LLM 理解的内容）
# ============================================================

@router.post("/chat/upload", summary="上传聊天附件并提取内容")
async def upload_chat_attachment(file: UploadFile = File(...)):
    """
    处理聊天消息中携带的附件文件，返回 LLM 可消费的内容。

    支持的文件类型：
    - **图片** (image/*): 转为 base64 data URL，供多模态模型直接使用
    - **PDF** (.pdf): 用 pdfplumber 提取全文，最多 30 000 字符
    - **Word 文档** (.docx): 用 python-docx 提取段落文本
    - **纯文本** (.txt / .md): 直接返回文本内容

    返回格式：
    ```json
    {"type": "image" | "text", "name": "文件名", "content": "...", "mime_type": "..."}
    ```
    """
    import io
    import base64

    filename: str = file.filename or "attachment"
    content_type: str = file.content_type or ""
    data: bytes = await file.read()

    ext = os.path.splitext(filename)[1].lower()

    # ── 图片 ──────────────────────────────────────────────────
    if content_type.startswith("image/") or ext in (".png", ".jpg", ".jpeg", ".gif", ".webp"):
        mime = content_type if content_type.startswith("image/") else f"image/{ext.lstrip('.')}"
        b64 = base64.b64encode(data).decode()
        data_url = f"data:{mime};base64,{b64}"
        return {"type": "image", "name": filename, "content": data_url, "mime_type": mime}

    # ── PDF ──────────────────────────────────────────────────
    if content_type == "application/pdf" or ext == ".pdf":
        try:
            import pdfplumber
            pages_text: List[str] = []
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        pages_text.append(t)
            text = "\n\n".join(pages_text)
        except Exception as e:
            logger.error(f"PDF extraction failed for '{filename}': {e}")
            raise HTTPException(status_code=422, detail=f"PDF 解析失败：{e}")
        _MAX_CHARS = 30_000
        if len(text) > _MAX_CHARS:
            text = text[:_MAX_CHARS] + "\n\n…（内容过长，已截断至前 30000 字符）"
        return {"type": "text", "name": filename, "content": text}

    # ── Word (.docx) ─────────────────────────────────────────
    if ext == ".docx" or content_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ):
        try:
            import docx as _docx
            doc = _docx.Document(io.BytesIO(data))
            paras = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n".join(paras)
        except ImportError:
            raise HTTPException(status_code=422, detail="DOCX 支持依赖 python-docx，请稍后重试（容器可能正在安装）")
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"DOCX 解析失败：{e}")
        _MAX_CHARS = 30_000
        if len(text) > _MAX_CHARS:
            text = text[:_MAX_CHARS] + "\n\n…（内容过长，已截断）"
        return {"type": "text", "name": filename, "content": text}

    # ── 纯文本 (.txt / .md) ──────────────────────────────────
    if content_type.startswith("text/") or ext in (".txt", ".md", ".csv", ".log"):
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            text = data.decode("latin-1", errors="replace")
        return {"type": "text", "name": filename, "content": text[:30_000]}

    raise HTTPException(
        status_code=415,
        detail=f"不支持的文件类型：{content_type or ext}。支持：image/*, .pdf, .docx, .txt, .md",
    )


@router.post("/intention", summary="意图分发接口")
async def intention_endpoint(request: IntentionRequest):
    """
    接收用户意图，Manager Agent 同步拆解任务并返回 task_ids。
    前端可据此轮询各任务完成状态，最终将结果回显在对话中。
    """
    result = await manager_agent.process_intention(request.intention)
    return result


# ============================================================
# 对话会话管理（Phase 7）
# ============================================================

@router.post("/conversations", status_code=201, summary="创建新会话")
async def create_conversation(request: ConversationCreate):
    """
    创建一个新的空会话。
    若提供 ``system_prompt``，自动插入一条 role=system 消息作为会话初始化。
    """
    from datetime import datetime as _dt
    conv = Conversation(
        title=request.title or "新对话",
        model=request.model,
    )
    with Session(engine) as session:
        session.add(conv)
        session.commit()
        session.refresh(conv)
        conv_id = conv.id

    # 可选 system prompt
    if request.system_prompt:
        with Session(engine) as session:
            db_msg = ConversationMessage(
                conversation_id=conv_id,
                role="system",
                content=request.system_prompt,
                token_count=_estimate_tokens(request.system_prompt),
                seq=0,
            )
            session.add(db_msg)
            conv_db = session.get(Conversation, conv_id)
            if conv_db:
                conv_db.message_count = 1
                conv_db.updated_at = _dt.utcnow()
                session.add(conv_db)
            session.commit()

    return {
        "id": conv_id,
        "title": request.title or "新对话",
        "model": request.model,
        "message_count": 1 if request.system_prompt else 0,
    }


@router.get("/conversations", summary="列出所有会话")
async def list_conversations(archived: bool = False):
    """
    返回会话摘要列表（按 updated_at 倒序，updated_at 为 null 时按 created_at 排序）。
    默认只返回未归档会话；``?archived=true`` 可查看已归档会话。
    """
    with Session(engine) as session:
        stmt = (
            select(Conversation)
            .where(Conversation.is_archived == archived)
        )
        convs = session.exec(stmt).all()

    # Python 侧排序（SQLite 对 NULL 的 ORDER BY 处理跨版本不稳定）
    convs_sorted = sorted(
        convs,
        key=lambda c: (c.updated_at or c.created_at or ""),
        reverse=True,
    )

    return {
        "conversations": [
            {
                "id": c.id,
                "title": c.title,
                "model": c.model,
                "message_count": c.message_count,
                "is_archived": c.is_archived,
                "created_at": c.created_at.isoformat() if c.created_at else None,
                "updated_at": c.updated_at.isoformat() if c.updated_at else None,
            }
            for c in convs_sorted
        ]
    }


@router.get("/conversations/{conv_id}", summary="获取会话详情（含全部消息）")
async def get_conversation(conv_id: str):
    """
    返回会话元数据 + 全部消息列表（按 seq 升序）。

    返回的消息中 ``tool_calls`` 字段已从 JSON 字符串反序列化为对象，
    前端可直接用于渲染工具调用卡片。
    """
    with Session(engine) as session:
        conv = session.get(Conversation, conv_id)
        if not conv:
            raise HTTPException(status_code=404, detail=f"Conversation '{conv_id}' not found.")
        msgs = session.exec(
            select(ConversationMessage)
            .where(ConversationMessage.conversation_id == conv_id)
            .order_by(ConversationMessage.seq.asc())
        ).all()

    return {
        "id": conv.id,
        "title": conv.title,
        "model": conv.model,
        "message_count": conv.message_count,
        "is_archived": conv.is_archived,
        "created_at": conv.created_at.isoformat() if conv.created_at else None,
        "updated_at": conv.updated_at.isoformat() if conv.updated_at else None,
        "messages": [
            {
                "role": m.role,
                "content": m.content,
                "tool_calls": json.loads(m.tool_calls_json) if m.tool_calls_json else None,
                "tool_call_id": m.tool_call_id,
                "seq": m.seq,
            }
            for m in msgs
        ],
    }


@router.patch("/conversations/{conv_id}", summary="更新会话属性")
async def update_conversation(conv_id: str, request: ConversationUpdate):
    """
    更新会话标题、模型或归档状态（PATCH 语义，仅修改提供的字段）。
    """
    from datetime import datetime as _dt
    with Session(engine) as session:
        conv = session.get(Conversation, conv_id)
        if not conv:
            raise HTTPException(status_code=404, detail=f"Conversation '{conv_id}' not found.")
        if request.title is not None:
            conv.title = request.title
        if request.model is not None:
            conv.model = request.model
        if request.is_archived is not None:
            conv.is_archived = request.is_archived
        conv.updated_at = _dt.utcnow()
        session.add(conv)
        session.commit()
    return {"status": "updated", "conversation_id": conv_id}


@router.delete("/conversations/{conv_id}", summary="删除会话")
async def delete_conversation(conv_id: str, hard: bool = False):
    """
    删除会话。

    - 默认软删除（标记 ``is_archived=True``），会话不再出现在列表中但可通过 ``?archived=true`` 查询。
    - ``?hard=true`` 时硬删除：同时删除会话下所有消息记录。
    """
    from datetime import datetime as _dt
    with Session(engine) as session:
        conv = session.get(Conversation, conv_id)
        if not conv:
            raise HTTPException(status_code=404, detail=f"Conversation '{conv_id}' not found.")
        if hard:
            msgs = session.exec(
                select(ConversationMessage).where(ConversationMessage.conversation_id == conv_id)
            ).all()
            for m in msgs:
                session.delete(m)
            session.delete(conv)
            action = "hard_deleted"
        else:
            conv.is_archived = True
            conv.updated_at = _dt.utcnow()
            session.add(conv)
            action = "archived"
        session.commit()
    return {"status": action, "conversation_id": conv_id}


@router.post("/conversations/{conv_id}/import", status_code=201, summary="批量导入历史消息（localStorage 迁移）")
async def import_messages(conv_id: str, request: ImportMessagesRequest):
    """
    一次性批量导入历史消息到指定会话（仅用于前端 localStorage → DB 迁移）。
    消息按传入顺序写入，seq 从当前最大值继续递增。
    迁移完成后此端点可保留，不影响正常功能。
    """
    from datetime import datetime as _dt
    with Session(engine) as session:
        conv = session.get(Conversation, conv_id)
        if not conv:
            raise HTTPException(status_code=404, detail=f"Conversation '{conv_id}' not found.")

        current_seq = conv.message_count
        for msg in request.messages:
            db_msg = ConversationMessage(
                conversation_id=conv_id,
                role=msg.role,
                content=msg.content,
                tool_call_id=getattr(msg, "tool_call_id", None),
                token_count=_estimate_tokens(msg.content or ""),
                seq=current_seq,
            )
            session.add(db_msg)
            current_seq += 1

        conv.message_count = current_seq
        conv.updated_at = _dt.utcnow()
        session.add(conv)
        session.commit()

    logger.info(f"import_messages: imported {len(request.messages)} messages into conv '{conv_id}'")
    return {"status": "imported", "conversation_id": conv_id, "imported_count": len(request.messages)}


# ============================================================
# 任务看板
# ============================================================

@router.get("/tasks", summary="查看所有任务状态")
async def get_tasks():
    """返回全量任务列表（含状态、依赖、时间戳、结果）。"""
    with Session(engine) as session:
        tasks = session.exec(select(Task)).all()
        return tasks


# ============================================================
# 技能库 CRUD
# ============================================================

@router.get("/skills", summary="获取技能库列表")
async def get_skills():
    """返回 SQLite 中注册的所有技能。"""
    with Session(engine) as session:
        skills = session.exec(select(SkillMetadata)).all()
        return {"skills": skills}


@router.post("/skills", summary="注册新技能", status_code=201)
async def create_skill(request: SkillCreate):
    """
    注册一个新技能：
    1. 写入 SQLite SkillMetadata 表（含 description 字段用于向量匹配）。
    2. 写入 ChromaDB skills_vector 集合。
    3. 若提供 script_content 则自动将内容写入 entrypoint 指定的文件。
    """
    # ---- 检查 ID 冲突 ----
    with Session(engine) as session:
        existing = session.exec(
            select(SkillMetadata).where(SkillMetadata.id == request.id)
        ).first()
        if existing:
            raise HTTPException(
                status_code=409, detail=f"Skill '{request.id}' already exists."
            )

    # ---- 写入脚本文件（可选）----
    if request.script_content:
        script_path = os.path.join(
            "/app" if os.path.exists("/app") else ".", request.entrypoint
        )
        os.makedirs(os.path.dirname(script_path), exist_ok=True)
        with open(script_path, "w", encoding="utf-8") as f:
            f.write(request.script_content)
        logger.info(f"Skill script written to {script_path}")

    # ---- Phase 8: 推断 error_tags（若未提供）----
    inferred_error_tags = request.error_tags or []
    if not inferred_error_tags:
        try:
            inferred_error_tags = await _infer_error_tags(
                skill_name=request.name,
                description=request.description,
                language=request.language,
                tags=[],
            )
        except Exception as e:
            logger.warning(f"error_tags inference failed (non-blocking): {e}")

    # ---- 写入 SQLite ----
    skill_meta = SkillMetadata(
        id=request.id,
        name=request.name,
        language=request.language,
        entrypoint=request.entrypoint,
        description=request.description,
        parameters_schema=request.parameters_schema,
        skill_type=request.skill_type,
        knowledge_content=request.knowledge_content,
        error_tags=inferred_error_tags,
    )
    with Session(engine) as session:
        session.add(skill_meta)
        session.commit()

    # ---- 写入 ChromaDB（via SkillManifest 兼容接口）----
    from app.models.schemas import SkillManifest
    skill_manifest = SkillManifest(
        id=request.id,
        name=request.name,
        description=request.description,
        language=request.language,
        content=request.script_content or "",
    )
    await memory_retriever.add_skill(skill_manifest)
    tool_registry.invalidate_cache()
    # Phase 11: 技能变更可能涉及脚本文件写入，触发增量索引更新
    _track_task(asyncio.create_task(code_indexer.update_incremental()))
    logger.info(f"Skill '{request.id}' registered successfully.")
    return {"status": "created", "skill_id": request.id}


# ---- 注意：以下固定路径路由必须在 /{skill_id} 参数路由之前 ----

@router.post("/skills/load-dir", summary="从 SKILL.md 目录批量导入技能", status_code=201)
async def load_skills_from_dir(skills_dir: str = "/app/skills"):
    """
    扫描 skills_dir 下所有含 SKILL.md 的子目录，解析并注册到系统中。
    遵循 Anthropic Agent Skills 协议文件夹格式：
        skill-name/
        ├── SKILL.md          # 必需，含 YAML frontmatter + Markdown 正文
        ├── scripts/          # 可选，可执行脚本
        └── resources/        # 可选，资源文件

    已注册的技能会跳过（不覆盖）；若要更新请先删除再重新导入。
    """
    parsed = skill_loader.load_dir(skills_dir)
    if not parsed:
        return {"status": "ok", "loaded": 0, "skipped": 0, "errors": [],
                "message": f"No SKILL.md found in {skills_dir}"}

    loaded, skipped, errors = 0, 0, []
    from app.models.schemas import SkillManifest
    for req in parsed:
        # 检查是否已存在
        with Session(engine) as session:
            if session.exec(select(SkillMetadata).where(SkillMetadata.id == req.id)).first():
                skipped += 1
                continue
        # 写入脚本文件
        if req.script_content and req.entrypoint:
            try:
                os.makedirs(os.path.dirname(req.entrypoint), exist_ok=True)
                with open(req.entrypoint, "w", encoding="utf-8") as f:
                    f.write(req.script_content)
            except OSError as e:
                errors.append(f"{req.id}: script write failed: {e}")
        # 写入 SQLite
        skill_meta = SkillMetadata(
            id=req.id, name=req.name, language=req.language,
            entrypoint=req.entrypoint, description=req.description,
            parameters_schema=req.parameters_schema,
        )
        with Session(engine) as session:
            session.add(skill_meta)
            session.commit()
        # 写入 ChromaDB
        skill_manifest = SkillManifest(
            id=req.id, name=req.name, description=req.description,
            language=req.language, content=req.script_content or "",
        )
        await memory_retriever.add_skill(skill_manifest)
        loaded += 1
        logger.info(f"Skill '{req.id}' imported from SKILL.md.")

    tool_registry.invalidate_cache()
    return {"status": "ok", "loaded": loaded, "skipped": skipped, "errors": errors}


@router.get("/skills/{skill_id}", summary="获取单个技能详情")
async def get_skill(skill_id: str):
    """返回 SQLite 中指定技能的完整元数据。"""
    with Session(engine) as session:
        skill = session.exec(
            select(SkillMetadata).where(SkillMetadata.id == skill_id)
        ).first()
    if not skill:
        raise HTTPException(status_code=404, detail=f"Skill '{skill_id}' not found.")
    return skill


@router.post("/skills/{skill_id}/run", summary="直接执行技能")
async def run_skill(skill_id: str, params: Optional[dict] = Body(default=None)):
    """
    直接调用指定技能的脚本并返回执行结果。

    请求体（可选 JSON 对象）会作为参数透传给脚本，格式应符合该技能
    `parameters_schema` 的定义。若不传则使用空参数 `{}`。

    返回：
        {"status": "success", "result": ..., "skill_id": ..., "language": ...}
        {"status": "error",   "message": ..., "skill_id": ...}
    """
    if params is None:
        params = {}

    with Session(engine) as session:
        skill = session.exec(
            select(SkillMetadata).where(SkillMetadata.id == skill_id)
        ).first()
    if not skill:
        raise HTTPException(status_code=404, detail=f"Skill '{skill_id}' not found.")

    exec_result = await skill_executor.run(
        language=skill.language,
        entrypoint=skill.entrypoint,
        params=params,
    )
    return {
        **exec_result,
        "skill_id": skill_id,
        "language": skill.language,
    }


@router.delete("/skills/{skill_id}", summary="删除技能")
async def delete_skill(skill_id: str):
    """从 SQLite 和 ChromaDB 中同步删除指定技能。"""
    with Session(engine) as session:
        skill = session.exec(
            select(SkillMetadata).where(SkillMetadata.id == skill_id)
        ).first()
        if not skill:
            raise HTTPException(status_code=404, detail=f"Skill '{skill_id}' not found.")
        session.delete(skill)
        session.commit()

    await memory_retriever.delete_skill(skill_id)
    tool_registry.invalidate_cache()
    # Phase 11: 触发增量索引更新
    _track_task(asyncio.create_task(code_indexer.update_incremental()))
    logger.info(f"Skill '{skill_id}' deleted.")
    return {"status": "deleted", "skill_id": skill_id}


@router.put("/skills/{skill_id}", summary="更新已有技能（PATCH 语义）")
async def update_skill(skill_id: str, request: SkillUpdate):
    """
    就地更新指定技能（Phase 4 — 技能自修正核心端点）。

    只更新请求体中明确提供的字段（None 字段不修改），
    同步更新 SQLite SkillMetadata 表 + ChromaDB skills_vector 集合。

    - `script_content` 非空时会覆盖 entrypoint 指向的脚本文件。
    - 适用场景：Agent 发现技能内容过时/不准确时自主调用修正。
    """
    with Session(engine) as session:
        skill = session.exec(
            select(SkillMetadata).where(SkillMetadata.id == skill_id)
        ).first()
        if not skill:
            raise HTTPException(status_code=404, detail=f"Skill '{skill_id}' not found.")

        # 更新提供的字段
        if request.name is not None:
            skill.name = request.name
        if request.description is not None:
            skill.description = request.description
        if request.language is not None:
            skill.language = request.language
        if request.entrypoint is not None:
            skill.entrypoint = request.entrypoint
        if request.parameters_schema is not None:
            skill.parameters_schema = request.parameters_schema
        if request.tags is not None:
            skill.tags = request.tags
        if request.skill_type is not None:
            skill.skill_type = request.skill_type
        if request.knowledge_content is not None:
            skill.knowledge_content = request.knowledge_content

        # 覆写脚本文件（可选）
        if request.script_content is not None:
            script_path = os.path.join(
                "/app" if os.path.exists("/app") else ".", skill.entrypoint
            )
            os.makedirs(os.path.dirname(script_path), exist_ok=True)
            with open(script_path, "w", encoding="utf-8") as f:
                f.write(request.script_content)
            logger.info(f"Skill '{skill_id}' script updated at {script_path}")

        session.add(skill)
        session.commit()
        session.refresh(skill)

    # 同步更新 ChromaDB（upsert）
    from app.models.schemas import SkillManifest
    skill_manifest = SkillManifest(
        id=skill_id,
        name=skill.name,
        description=skill.description,
        language=skill.language,
        content=request.script_content or "",
    )
    await memory_retriever.add_skill(skill_manifest)
    tool_registry.invalidate_cache()
    # Phase 11: 触发增量索引更新
    _track_task(asyncio.create_task(code_indexer.update_incremental()))
    logger.info(f"Skill '{skill_id}' updated.")
    return {"status": "updated", "skill_id": skill_id}


# ============================================================
# 代理 / 模型 / 知识库 / 错题集
# ============================================================

@router.get("/proxy/status", summary="查看代理健康状态")
async def get_proxy_status():
    """返回当前代理池的连接健康状态。"""
    return await proxy_manager.get_health_status()


@router.get("/models", summary="获取可用模型列表")
async def get_models():
    return {
        "available_models": model_router.available_models,
        "default_model": model_router.default_model,
    }


# ============================================================
# 运行时设置（Phase 9）
# ============================================================

# 设置项默认值与可选项定义
_SETTINGS_DEFAULTS = {
    "vision_model": {
        "value": "gemini-2.5-flash",
        "description": "PDF 图片识别使用的多模态模型",
        "options": [],  # 动态从 model_router 获取支持 vision 的模型
    },
    "max_tool_rounds": {
        "value": "30",
        "description": "Tool Calling 最大循环轮次（防止 LLM 无限循环，推荐范围：10-50）",
        "options": [],  # 自由填写整数
    },
}

# 内置回退列表（当 API 查询失败或尚未刷新时使用；仅 Gemini 侧）
_VISION_CAPABLE_MODELS_FALLBACK: List[str] = [
    "gemini-2.5-pro",
    "gemini-2.5-flash",
    "gemini-2.5-flash-lite",
    "gemini-2.0-flash",
]

# 动态 Vision 模型缓存，应用启动后默认使用内置回退列表
_vision_models_cache: Dict[str, Any] = {
    "models": _VISION_CAPABLE_MODELS_FALLBACK.copy(),
    "updated_at": None,  # None 表示从未从 Gemini API 刷新过
}


async def _fetch_vision_models_from_api() -> List[str]:
    """
    从 Gemini REST API 和 Volcengine ARK API 检测支持图片输入（Vision）的模型。

    检测策略（按优先级）：
    1. Gemini：查询 /v1beta/models，检查 inputModalities 字段；
       若字段不存在则基于模型名称推断（所有 flash/pro/2.0+ 均支持 vision）。
    2. Volcengine：查询 ARK API /api/v3/models，检查 input_types / input_modalities 字段；
       若 API 不返回能力信息则基于模型名称中的 vision 关键词推断。

    最终结果与 model_router.available_models 取交集，只返回系统已配置可用的模型。
    通过 Clash 代理（settings.proxy_url）访问 Google API。
    """
    from app.core.config import settings as _app_settings

    vision_models: List[str] = []

    # ── 1. Gemini Vision 检测 ────────────────────────────────────────────
    gemini_known: set = set(model_router.available_models.get("gemini", []))
    try:
        url = (
            f"https://generativelanguage.googleapis.com/v1beta/models"
            f"?key={_app_settings.gemini_api_key}&pageSize=200"
        )
        async with httpx.AsyncClient(
            proxy=_app_settings.proxy_url,
            trust_env=False,
            timeout=30,
        ) as client:
            resp = await client.get(url)
            resp.raise_for_status()
            data = resp.json()

        for model in data.get("models", []):
            raw_name: str = model.get("name", "")
            short_name = raw_name.replace("models/", "")
            methods: List[str] = model.get("supportedGenerationMethods", [])
            input_modalities: List[str] = model.get("inputModalities", [])

            # 必须支持 generateContent（排除 embed / aqa 等专用模型）
            if "generateContent" not in methods:
                continue
            # 只保留 available_models 中已配置的 Gemini 模型
            if short_name not in gemini_known:
                continue

            # 检测 Vision 能力
            if input_modalities:
                # 优先使用 API 返回的 inputModalities 字段（新版 Gemini API）
                has_vision = any("image" in m.lower() for m in input_modalities)
            else:
                # 回退：gemini-2.0 / 2.5 / 3.x 的 flash / pro 系列均支持多模态视觉输入
                has_vision = any(
                    kw in short_name
                    for kw in ("flash", "pro", "gemini-2.", "gemini-3.")
                )

            if has_vision:
                vision_models.append(short_name)

        logger.info(
            f"Vision refresh – Gemini: {len([m for m in vision_models if 'gemini' in m])} "
            f"model(s) detected from API"
        )
    except Exception as e:
        logger.warning(f"Gemini vision check via API failed: {e}. Falling back to known list.")
        # 回退：将 available_models 中所有 Gemini 模型视为支持 vision（排除纯文本/嵌入模型）
        for m in gemini_known:
            if not any(skip in m for skip in ("text-", "embedding-", "aqa")):
                vision_models.append(m)

    # ── 2. Volcengine Vision 检测 ─────────────────────────────────────────
    ve_known: set = set(model_router.available_models.get("volcengine", []))
    try:
        ve_api_url = "https://ark.cn-beijing.volces.com/api/v3/models"
        headers = {"Authorization": f"Bearer {_app_settings.volcengine_api_key}"}
        async with httpx.AsyncClient(trust_env=False, timeout=20) as client:
            resp = await client.get(ve_api_url, headers=headers)
            resp.raise_for_status()
            ve_data = resp.json()

        # id -> model_info 映射（ARK OpenAI-compat API）
        ve_api_map: Dict[str, Any] = {
            m.get("id", ""): m for m in ve_data.get("data", [])
        }

        api_hit = 0
        for model_id in ve_known:
            model_info = ve_api_map.get(model_id, {})
            # ARK 平台扩展字段：input_types / input_modalities / supported_input_types
            input_types: List[str] = (
                model_info.get("input_types")
                or model_info.get("input_modalities")
                or model_info.get("supported_input_types")
                or []
            )
            if input_types:
                api_hit += 1
                if any("image" in t.lower() for t in input_types):
                    vision_models.append(model_id)
            else:
                # API 未返回能力信息：通过名称推断
                if "vision" in model_id.lower():
                    vision_models.append(model_id)

        logger.info(
            f"Vision refresh – Volcengine: capability info from API for "
            f"{api_hit}/{len(ve_known)} model(s)"
        )
    except Exception as e:
        logger.warning(f"Volcengine vision check via API failed: {e}. Using name-based inference.")
        # 回退：仅通过模型名称中的 vision 关键词识别
        for m in ve_known:
            if "vision" in m.lower():
                vision_models.append(m)

    # 新版本 / 字母序靠前的排在前面
    return sorted(set(vision_models), reverse=True)


@router.post("/settings/vision-models/refresh", summary="刷新 Vision 模型列表")
async def refresh_vision_models():
    """
    向 Gemini REST API 查询当前所有可用模型，筛选支持多模态图片输入的模型，
    更新内存缓存并返回最新列表及刷新时间。
    """
    global _vision_models_cache
    try:
        models = await _fetch_vision_models_from_api()
        updated_at = _dt.utcnow().strftime("%Y-%m-%d %H:%M UTC")
        _vision_models_cache = {"models": models, "updated_at": updated_at}
        logger.info(f"Vision models refreshed from API: {len(models)} models, updated_at={updated_at}")
        return {
            "models": models,
            "updated_at": updated_at,
            "count": len(models),
        }
    except Exception as e:
        logger.error(f"Vision models refresh failed: {e}")
        raise HTTPException(status_code=502, detail=f"Gemini API 查询失败：{e}")


@router.get("/settings", summary="获取所有运行时设置")
async def get_settings():
    """
    返回所有运行时设置项的当前值、描述及可选项。

    每个设置项包含：
    - `value`       — 当前生效值
    - `description` — 可读描述
    - `options`     — 可选值列表（空列表表示自由填写）
    """
    from app.models.database import RuntimeSetting

    settings_out = {}
    for key, defaults in _SETTINGS_DEFAULTS.items():
        with Session(engine) as session:
            db_setting = session.get(RuntimeSetting, key)
        current_value = db_setting.value if db_setting else defaults["value"]

        # 动态构建 vision 模型可选项
        options = defaults.get("options", [])
        extra: Dict[str, Any] = {}
        if key == "vision_model":
            options = _vision_models_cache["models"]
            extra["last_updated"] = _vision_models_cache["updated_at"] or "（使用内置列表，点击刷新获取最新）"

        settings_out[key] = {
            "value": current_value,
            "description": defaults["description"],
            "options": options,
            **extra,
        }

    return {"settings": settings_out}


@router.put("/settings/{key}", summary="更新单个运行时设置")
async def update_setting(key: str, body: dict = Body(...)):
    """
    更新指定 key 的运行时设置值。

    Request Body::

        {"value": "gemini-2.5-flash"}
    """
    from app.models.database import RuntimeSetting
    from datetime import datetime as _dt

    if key not in _SETTINGS_DEFAULTS:
        raise HTTPException(status_code=400, detail=f"未知的设置项: '{key}'")

    new_value = body.get("value")
    if new_value is None:
        raise HTTPException(status_code=400, detail="请求体必须包含 'value' 字段。")

    # max_tool_rounds 需要是 1-100 之间的整数
    if key == "max_tool_rounds":
        try:
            rounds_val = int(new_value)
            if rounds_val < 1 or rounds_val > 100:
                raise HTTPException(status_code=400, detail="max_tool_rounds 必须在 1 到 100 之间。")
        except (ValueError, TypeError):
            raise HTTPException(status_code=400, detail="max_tool_rounds 必须是整数。")

    with Session(engine) as session:
        existing = session.get(RuntimeSetting, key)
        if existing:
            existing.value = str(new_value)
            existing.updated_at = _dt.utcnow()
            session.add(existing)
        else:
            session.add(RuntimeSetting(
                key=key,
                value=str(new_value),
                description=_SETTINGS_DEFAULTS[key]["description"],
            ))
        session.commit()

    logger.info(f"Setting '{key}' updated to '{new_value}'")
    return {"key": key, "value": str(new_value), "status": "ok"}


def get_runtime_setting(key: str) -> str:
    """
    读取运行时设置的当前值（供其他服务模块调用）。

    优先从 DB 读取；DB 无记录时返回默认值。
    该函数为同步调用，适合在非 async 上下文中使用。
    """
    from app.models.database import RuntimeSetting

    with Session(engine) as session:
        setting = session.get(RuntimeSetting, key)
    if setting:
        return setting.value
    return _SETTINGS_DEFAULTS.get(key, {}).get("value", "")


@router.get("/knowledge", summary="获取知识库列表")
async def get_knowledge():
    """返回根目录下的 Markdown 文档作为知识库（不含 error_ledger.md）。"""
    root_dir = "/app" if os.path.exists("/app/features.md") else "."
    knowledge_files = [
        f for f in os.listdir(root_dir)
        if f.endswith(".md") and f != "error_ledger.md"
    ]
    return {"knowledge_files": knowledge_files}


# ============================================================
# Phase 8 — 标签化错题集 CRUD
# ============================================================

async def _auto_generate_error_entry(task_description: str, error_msg: str) -> dict:
    """
    调用 LLM 为失败任务自动生成结构化错题（含标签）。
    输出 JSON：{"title", "context", "correction", "prevention", "tags", "severity"}
    """
    import re as _re
    prompt = (
        f"你是一个错误分析助手。请分析以下任务失败信息，生成一条结构化的错题记录。\n\n"
        f"任务描述：{task_description}\n错误信息：{error_msg}\n\n"
        '请以 JSON 格式输出（不要输出其他内容）：\n'
        '{\n'
        '    "title": "简短标题（10-20字概括错误）",\n'
        '    "context": "详细的错误上下文描述",\n'
        '    "correction": "正确的解决方案",\n'
        '    "prevention": "预防再次发生的建议",\n'
        '    "tags": ["3-5个分类标签，如 python, docker, api, encoding, timeout 等"],\n'
        '    "severity": "critical 或 warning 或 info"\n'
        '}'
    )
    try:
        response = await model_router.generate(
            messages=[Message(role="user", content=prompt)],
            temperature=0.3,
            max_tokens=500,
        )
        text = response.content or ""
        json_match = _re.search(r'\{.*\}', text, _re.DOTALL)
        if json_match:
            return json.loads(json_match.group())
    except Exception as e:
        logger.warning(f"_auto_generate_error_entry LLM call failed: {e}")

    # fallback
    return {
        "title": (task_description or "未知任务")[:30],
        "context": f"任务描述: {task_description}\n错误信息: {error_msg}",
        "correction": "请检查任务依赖的技能是否存在，或调整任务描述后重新提交。",
        "prevention": "",
        "tags": ["general"],
        "severity": "warning",
    }


async def _infer_error_tags(skill_name: str, description: str, language: str, tags: list) -> list:
    """调用 LLM 推断技能关联的错题标签。"""
    import re as _re
    prompt = (
        f"你是一个软件工程助手。请根据以下技能信息，推断执行此技能时可能遇到的错误类别。\n\n"
        f"技能名称：{skill_name}\n技能描述：{description}\n语言：{language}\n"
        f"标签：{', '.join(tags) if tags else '无'}\n\n"
        "请只输出一个 JSON 数组，包含 3-8 个错误类别标签，例如：\n"
        '["python", "file-io", "encoding", "timeout", "dependency"]\n\n'
        "不要输出其他内容。"
    )
    try:
        response = await model_router.generate(
            messages=[Message(role="user", content=prompt)],
            temperature=0.2,
            max_tokens=100,
        )
        text = response.content or ""
        arr_match = _re.search(r'\[.*\]', text, _re.DOTALL)
        if arr_match:
            result = json.loads(arr_match.group())
            if isinstance(result, list):
                return result
    except Exception as e:
        logger.warning(f"_infer_error_tags LLM call failed: {e}")
    # fallback
    return list(set((tags or []) + [language]))


@router.get("/errors", summary="获取错题集（兼容旧接口）")
async def get_errors():
    """保留旧接口向后兼容：返回 error_ledger.md 文件内容。"""
    error_file = (
        "/app/error_ledger.md"
        if os.path.exists("/app/error_ledger.md")
        else "error_ledger.md"
    )
    if os.path.exists(error_file):
        with open(error_file, "r", encoding="utf-8") as f:
            return {"content": f.read()}
    return {"content": "暂无错题集记录。"}


@router.post("/errors/ingest", summary="将 error_ledger.md 重新写入向量库（兼容旧接口）")
async def ingest_errors():
    """保留旧接口向后兼容：解析 error_ledger.md 写入 ChromaDB。"""
    error_file = (
        "/app/error_ledger.md"
        if os.path.exists("/app/error_ledger.md")
        else "error_ledger.md"
    )
    count = await memory_retriever.ingest_error_ledger(error_file)
    return {"status": "ok", "entries_ingested": count}


@router.post("/errors/entries", summary="创建错题条目", status_code=201)
async def create_error_entry(request: ErrorEntryCreate):
    """创建一条结构化错题。写入 SQLite（权威源），后台自动同步到 ChromaDB。"""
    entry = ErrorEntry(
        title=request.title,
        context=request.context,
        correction=request.correction,
        prevention=request.prevention,
        tags=request.tags,
        severity=request.severity,
        source=request.source,
        related_skill_ids=request.related_skill_ids,
    )
    with Session(engine) as session:
        session.add(entry)
        session.commit()
        session.refresh(entry)
        entry_id = entry.id

    # 后台同步到 ChromaDB（fire-and-forget，不阻塞响应）
    asyncio.ensure_future(vector_sync.sync_error_entry(
        entry_id=entry_id,
        context=request.context,
        correction=request.correction,
        tags=request.tags,
        severity=request.severity,
    ))
    return {"id": entry_id, "title": request.title, "tags": request.tags}


@router.get("/errors/entries", summary="列出错题条目（支持标签筛选 + 关键词搜索）")
async def list_error_entries(
    tags: Optional[str] = None,
    search: Optional[str] = None,
    severity: Optional[str] = None,
    limit: int = 50,
    offset: int = 0,
):
    """
    返回错题列表。

    - ``?tags=python,docker`` — 按标签筛选（OR 逻辑，命中任一标签即返回）
    - ``?search=关键词`` — 全文模糊搜索 title / context
    - ``?severity=critical`` — 按严重程度筛选
    - 默认按 hit_count DESC + created_at DESC 排序
    """
    # 在 Session 内完成所有数据读取，避免 detached instance 问题
    with Session(engine) as session:
        entries = session.exec(select(ErrorEntry)).all()

        result = list(entries)

        if tags:
            tag_list = [t.strip() for t in tags.split(",") if t.strip()]
            result = [e for e in result if any(t in (e.tags or []) for t in tag_list)]

        if severity:
            result = [e for e in result if e.severity == severity]

        if search:
            kw = search.lower()
            result = [e for e in result if kw in (e.title or "").lower() or kw in (e.context or "").lower()]

        result.sort(key=lambda e: (e.hit_count or 0, str(e.created_at or "")), reverse=True)

        total = len(result)
        result = result[offset: offset + limit]

        # 在 Session 内序列化为 dict，确保所有字段（包括 JSON 列）已加载
        serialized = [
            {
                "id": e.id,
                "title": e.title,
                "tags": e.tags or [],
                "severity": e.severity,
                "source": e.source,
                "hit_count": e.hit_count or 0,
                "created_at": e.created_at.isoformat() if e.created_at else None,
            }
            for e in result
        ]

    return {
        "total": total,
        "entries": serialized,
    }


@router.get("/errors/tags", summary="列出所有错题标签及频次")
async def list_error_tags():
    """返回所有标签及其出现次数，供前端标签云/筛选器使用。"""
    with Session(engine) as session:
        entries = session.exec(select(ErrorEntry)).all()
        tag_count: dict = {}
        for e in entries:
            for t in (e.tags or []):
                tag_count[t] = tag_count.get(t, 0) + 1
    sorted_tags = sorted(tag_count.items(), key=lambda x: x[1], reverse=True)
    return {"tags": [{"name": t, "count": c} for t, c in sorted_tags]}


@router.get("/errors/entries/{entry_id}", summary="获取错题详情")
async def get_error_entry(entry_id: str):
    with Session(engine) as session:
        entry = session.get(ErrorEntry, entry_id)
        if not entry:
            raise HTTPException(404, f"ErrorEntry '{entry_id}' not found.")
        # 在 Session 内读取所有字段，避免 detached instance 问题
        data = {
            "id": entry.id,
            "title": entry.title,
            "context": entry.context,
            "correction": entry.correction,
            "prevention": entry.prevention or "",
            "tags": entry.tags or [],
            "severity": entry.severity,
            "source": entry.source,
            "related_skill_ids": entry.related_skill_ids or [],
            "hit_count": entry.hit_count or 0,
            "created_at": entry.created_at.isoformat() if entry.created_at else None,
        }
    return data


@router.delete("/errors/entries/{entry_id}", summary="删除错题")
async def delete_error_entry(entry_id: str):
    with Session(engine) as session:
        entry = session.get(ErrorEntry, entry_id)
        if not entry:
            raise HTTPException(404, f"ErrorEntry '{entry_id}' not found.")
        session.delete(entry)
        session.commit()
    # 后台同步删除 ChromaDB 副本
    asyncio.ensure_future(vector_sync.remove_error_entry(entry_id))
    return {"status": "deleted", "id": entry_id}


@router.post("/errors/migrate", summary="从 error_ledger.md 迁移旧数据到 SQLite（直接解析，不依赖 LLM）")
async def migrate_error_ledger(force: bool = False):
    """
    解析 error_ledger.md 的结构化 Markdown，直接提取字段写入 ErrorEntry SQLite 表 + ChromaDB。
    不依赖 LLM，保证内容原文不变，避免编码丢失。
    幂等：已存在的条目按 title MD5 跳过。

    - ``?force=true`` — 强制覆盖已存在的条目（修复历史乱码数据）
    """
    import re as _re
    import hashlib as _hashlib

    error_file = (
        "/app/error_ledger.md"
        if os.path.exists("/app/error_ledger.md")
        else "error_ledger.md"
    )
    if not os.path.exists(error_file):
        return {"status": "skipped", "message": "error_ledger.md not found"}

    with open(error_file, "r", encoding="utf-8") as f:
        content = f.read()

    sections = _re.split(r"\n### ", content)
    migrated, skipped = 0, 0

    # 提取 ## 级别分类标题，用于推断 tags
    category_map: Dict[int, str] = {}  # line_offset -> category name
    for m in _re.finditer(r"^## \[[\d-]+\]\s*(.+)$", content, _re.MULTILINE):
        category_map[m.start()] = m.group(1).strip()

    # 常见关键词 → 标签映射（用于自动打标签）
    _TAG_KEYWORDS = {
        "docker": ["docker", "container", "容器", "dockerfile", "docker-compose", "镜像"],
        "python": ["python", "pip", "pydantic", "numpy", "module", "importerror", "syntaxerror"],
        "network": ["proxy", "代理", "网络", "http", "api", "502", "timeout", "超时"],
        "database": ["sqlite", "chromadb", "数据库", "db", "unique", "constraint"],
        "encoding": ["编码", "utf-8", "encoding", "乱码", "charset"],
        "deployment": ["部署", "端口", "port", "deploy", "构建", "build"],
        "asyncio": ["asyncio", "queue", "队列", "worker", "任务"],
        "proxy": ["clash", "mihomo", "shadowsocks", "ss-2022", "订阅"],
        "config": ["配置", "config", "env", "环境变量", "pydantic settings"],
        "pdf": ["pdf", "pypdf", "pdfplumber"],
        "llm": ["llm", "gemini", "模型", "model", "token"],
    }

    def _infer_tags_from_text(text: str) -> list:
        """从文本中推断标签（纯本地，不调用 LLM）。"""
        text_lower = text.lower()
        tags = set()
        for tag, keywords in _TAG_KEYWORDS.items():
            for kw in keywords:
                if kw.lower() in text_lower:
                    tags.add(tag)
                    break
        return sorted(tags) if tags else ["general"]

    def _parse_md_entry(raw: str) -> dict:
        """
        解析单条 ### 级 MD 条目，提取结构化字段。

        支持的格式:
            ### N. Title
            - **问题描述**: ...
            - **原因分析**: ...
            - **解决方案**: ...
            - **预防建议**: ...
        """
        lines = raw.strip().split("\n")
        title = _re.sub(r"^\d+\.\s*", "", lines[0].strip())  # 去掉序号前缀
        body = "\n".join(lines[1:]).strip()

        # 提取各 **字段**: 值
        field_pattern = _re.compile(r"-\s*\*\*(.+?)\*\*\s*[:：]\s*(.*?)(?=\n-\s*\*\*|\Z)", _re.DOTALL)
        fields: Dict[str, str] = {}
        for m in field_pattern.finditer(body):
            key = m.group(1).strip()
            val = m.group(2).strip()
            fields[key] = val

        # 组装 context：问题描述 + 发现途径 + 原因分析
        context_parts = []
        for k in ["问题描述", "发现途径", "原因分析"]:
            if k in fields:
                context_parts.append(f"{k}: {fields[k]}")
        context = "\n".join(context_parts) if context_parts else body

        correction = fields.get("解决方案", "")
        prevention = fields.get("预防建议", "")

        return {
            "title": title,
            "context": context,
            "correction": correction,
            "prevention": prevention,
        }

    for raw in sections[1:]:
        lines = raw.strip().split("\n")
        raw_title = lines[0].strip()
        body = "\n".join(lines[1:]).strip()
        if not raw_title or not body:
            continue

        doc_id = _hashlib.md5(raw_title.encode("utf-8")).hexdigest()
        with Session(engine) as session:
            existing = session.get(ErrorEntry, doc_id)
            if existing and not force:
                skipped += 1
                continue
            # force 模式：删除旧数据以便重写
            if existing and force:
                session.delete(existing)
                session.commit()

        parsed = _parse_md_entry(raw)
        tags = _infer_tags_from_text(raw_title + " " + body)

        entry = ErrorEntry(
            id=doc_id,
            title=parsed["title"],
            context=parsed["context"],
            correction=parsed["correction"],
            prevention=parsed["prevention"],
            tags=tags,
            severity="info",
            source="manual",
        )
        with Session(engine) as session:
            session.add(entry)
            session.commit()

        migrated += 1

    # 迁移完成后，一次性全量重建 ChromaDB（比逐条同步高效）
    if migrated > 0:
        asyncio.ensure_future(vector_sync.full_rebuild_errors())

    return {"status": "ok", "migrated": migrated, "skipped": skipped}


@router.post("/errors/sync", summary="全量重建 ChromaDB 错题向量索引")
async def sync_error_vectors():
    """
    从 SQLite ErrorEntry 表全量重建 ChromaDB error_ledger_vector。

    适用场景：
    - 发现前端能看到但 Worker RAG 搜不到的错题
    - ChromaDB 数据损坏或被意外清空
    - 任何数据不一致的情况
    """
    count = await vector_sync.full_rebuild_errors()
    return {"status": "ok", "synced": count}


@router.post("/skills/sync", summary="全量重建 ChromaDB 技能向量索引")
async def sync_skill_vectors():
    """从 SQLite SkillMetadata 表全量重建 ChromaDB skills_vector。"""
    count = await vector_sync.full_rebuild_skills()
    return {"status": "ok", "synced": count}


# Phase 4 — PDF-to-Skills 流水线
# ============================================================

_PDF_UPLOAD_DIR = "/app/data/pdfs"
_MAX_UPLOAD_SIZE = 250 * 1024 * 1024   # 250MB
_UPLOAD_CHUNK_SIZE = 1024 * 1024       # 1MB 分块写盘


@router.post("/pdf/upload", summary="上传 PDF 文件", status_code=201)
async def upload_pdf(file: UploadFile = File(...)):
    """
    上传 PDF 文件到服务器，返回 doc_id 供后续流水线调用。

    - 自动计算 SHA-256 哈希，同一文件不重复入库（幂等）。
    - 使用 multipart/form-data 上传；前端用 `<input type="file">` 或 curl `-F` 发送。
    - 文件大小限制 250MB；超限立即中断，返回 HTTP 413。
    - 流式分块写盘，内存占用固定 ~1MB。

    Response::

        {
            "doc_id": "uuid",
            "filename": "example.pdf",
            "file_path": "/app/data/pdfs/uuid.pdf",
            "page_count": 320,
            "file_hash": "sha256..."
        }
    """
    from app.services.pdf_processor import pdf_processor

    if not file.filename or not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="只支持 PDF 文件（.pdf 后缀）。")

    os.makedirs(_PDF_UPLOAD_DIR, exist_ok=True)

    # 临时写入磁盘（流式分块，内存固定）
    tmp_path = os.path.join(_PDF_UPLOAD_DIR, f"tmp_{uuid.uuid4().hex}.pdf")
    total_size = 0
    try:
        with open(tmp_path, "wb") as f:
            while True:
                chunk = await file.read(_UPLOAD_CHUNK_SIZE)
                if not chunk:
                    break
                total_size += len(chunk)
                if total_size > _MAX_UPLOAD_SIZE:
                    # 超限：立即删除临时文件，返回 413
                    f.close()
                    os.remove(tmp_path)
                    raise HTTPException(
                        status_code=413,
                        detail=f"文件大小超过 {_MAX_UPLOAD_SIZE // (1024*1024)}MB 上限。",
                    )
                f.write(chunk)

        file_hash = pdf_processor.compute_file_hash(tmp_path)

        # 幂等检查：同一哈希是否已上传
        with Session(engine) as session:
            existing = session.exec(
                select(PDFDocument).where(PDFDocument.file_hash == file_hash)
            ).first()
        if existing:
            os.remove(tmp_path)
            logger.info(f"PDF already uploaded: {existing.id}")
            return {
                "doc_id": existing.id,
                "filename": existing.filename,
                "file_path": existing.file_path,
                "page_count": existing.page_count,
                "file_hash": existing.file_hash,
                "note": "File already uploaded, returning existing doc_id.",
            }

        # 正式存储
        doc_id = str(uuid.uuid4())
        final_path = os.path.join(_PDF_UPLOAD_DIR, f"{doc_id}.pdf")
        os.rename(tmp_path, final_path)

        # 快速获取页数（不阻塞）
        import asyncio as _asyncio
        loop = _asyncio.get_event_loop()
        page_count = await loop.run_in_executor(None, _count_pdf_pages, final_path)

        # 写入 PDFDocument
        doc = PDFDocument(
            id=doc_id,
            filename=file.filename,
            file_path=final_path,
            file_hash=file_hash,
            page_count=page_count,
            status="pending",
        )
        with Session(engine) as session:
            session.add(doc)
            session.commit()

        logger.info(f"PDF uploaded: {doc_id} ({file.filename}, {page_count} pages)")
        return {
            "doc_id": doc_id,
            "filename": file.filename,
            "file_path": final_path,
            "page_count": page_count,
            "file_hash": file_hash,
        }
    except Exception as e:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
        logger.error(f"PDF upload failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


def _count_pdf_pages(path: str) -> int:
    """同步快速统计 PDF 页数（run_in_executor 调用）。"""
    try:
        import pypdf
        with pypdf.PdfReader(path) as r:
            return len(r.pages)
    except Exception:
        return 0


@router.post("/pdf/to-skills", summary="触发 PDF→Skills 异步流水线")
async def pdf_to_skills(request: PDFToSkillsRequest):
    """
    触发 PDF-to-Skills 全流水线（异步，立即返回）。

    参数二选一：
    - `doc_id`  — 来自 `/pdf/upload` 返回的 doc_id（推荐）
    - `pdf_path` — 容器内 PDF 文件路径（直接指定）

    流水线在后台执行，通过 `GET /pdf/status/{doc_id}` 查询进度。

    Response::

        {"status": "accepted", "doc_id": "uuid"}
    """
    from app.services.pdf_processor import pdf_processor

    pdf_path: Optional[str] = None
    doc_id: Optional[str] = request.doc_id

    if doc_id:
        with Session(engine) as session:
            doc = session.exec(
                select(PDFDocument).where(PDFDocument.id == doc_id)
            ).first()
        if not doc:
            raise HTTPException(status_code=404, detail=f"PDFDocument '{doc_id}' not found.")
        pdf_path = doc.file_path
    elif request.pdf_path:
        pdf_path = request.pdf_path
        # 如果没有 doc_id，创建一个新的 PDFDocument 记录
        file_hash = pdf_processor.compute_file_hash(pdf_path)
        doc_id = str(uuid.uuid4())
        page_count = _count_pdf_pages(pdf_path)
        doc = PDFDocument(
            id=doc_id,
            filename=os.path.basename(pdf_path),
            file_path=pdf_path,
            file_hash=file_hash,
            page_count=page_count,
            status="pending",
        )
        with Session(engine) as session:
            session.add(doc)
            session.commit()
    else:
        raise HTTPException(status_code=400, detail="必须提供 doc_id 或 pdf_path 之一。")

    if not pdf_path or not os.path.exists(pdf_path):
        raise HTTPException(status_code=404, detail=f"PDF 文件不存在: {pdf_path}")

    # 标记为处理中
    with Session(engine) as session:
        doc = session.exec(select(PDFDocument).where(PDFDocument.id == doc_id)).first()
        if doc:
            doc.status = "processing"
            session.add(doc)
            session.commit()

    # 后台异步执行流水线
    async def _run_pipeline():
        try:
            result = await pdf_processor.pdf_to_skills(
                pdf_path=pdf_path,
                output_dir=request.output_dir,
                skill_prefix=request.skill_prefix,
                max_tokens_per_chunk=request.max_tokens_per_chunk,
                doc_id=doc_id,
            )
            logger.info(
                f"Pipeline completed for doc {doc_id}: "
                f"{result.skills_registered} skills registered."
            )
        except Exception as e:
            logger.error(f"Pipeline failed for doc {doc_id}: {e}")
            with Session(engine) as session:
                d = session.exec(select(PDFDocument).where(PDFDocument.id == doc_id)).first()
                if d:
                    d.status = "failed"
                    d.error_msg = str(e)
                    session.add(d)
                    session.commit()

    _track_task(asyncio.create_task(_run_pipeline()))
    return {"status": "accepted", "doc_id": doc_id}


@router.get("/pdf/status/{doc_id}", summary="查询 PDF 流水线处理进度")
async def get_pdf_status(doc_id: str):
    """
    查询指定 PDFDocument 的流水线处理状态与结果摘要。

    Status 值说明：
    - `pending`    — 已上传，等待触发
    - `processing` — 流水线运行中
    - `completed`  — 已完成，可查看 skills_generated
    - `failed`     — 流水线失败，见 error_msg
    
    Phase 9 增强（B-4）：新增 processed_chunks 和 progress_pct 字段用于进度上报。
    """
    with Session(engine) as session:
        doc = session.exec(
            select(PDFDocument).where(PDFDocument.id == doc_id)
        ).first()
    if not doc:
        raise HTTPException(status_code=404, detail=f"PDFDocument '{doc_id}' not found.")
    
    # Phase 9 B-4: 计算进度百分比
    progress_pct = 0.0
    if doc.total_chunks and doc.total_chunks > 0:
        progress_pct = round(doc.processed_chunks / doc.total_chunks * 100, 1)
    
    return {
        "doc_id": doc.id,
        "status": doc.status,
        "file_name": doc.file_name,
        "total_chunks": doc.total_chunks,
        "processed_chunks": doc.processed_chunks,  # Phase 9 B-4
        "progress_pct": progress_pct,             # Phase 9 B-4
        "skills_generated": doc.skills_generated or [],
        "error_msg": doc.error_msg,
        "created_at": doc.created_at,
        "completed_at": doc.completed_at,
    }


@router.get("/pdf", summary="列出所有已上传的 PDF 文档")
async def list_pdfs():
    """返回所有 PDFDocument 记录（用于前端管理界面）。"""
    with Session(engine) as session:
        docs = session.exec(select(PDFDocument)).all()
    return {"documents": docs}


# ============================================================
# Phase 5 — ms-agent 深度能力
# ============================================================

@router.post(
    "/research/deep",
    response_model=DeepResearchResponse,
    summary="[ms-agent] 触发 deep_research v2 深度研究任务",
    tags=["msagent"],
)
async def start_deep_research(request: DeepResearchRequest):
    """
    启动一个 **deep_research v2** 深度研究任务（异步 fire-and-forget）。

    工作流（ms-agent 内部）：
    1. ResearcherAgent 制定研究计划
    2. SearcherAgent 并行检索（arxiv / EXA / SerpAPI）
    3. ReporterAgent 汇总证据，生成结构化报告

    产物文件（容器内 `/app/data/outputs/research/{task_id}/`）：
    - `final_report.md`  — 最终研究报告（Markdown）
    - `plan.json`        — 研究计划
    - `evidence/`        — 原始证据片段
    - `reports/`         — 中间分析报告

    立即返回 `task_id`，通过 `GET /research/{task_id}` 轮询状态和读取报告。

    > 注意：任务运行时间通常 5-30 分钟，取决于研究深度和搜索引擎可用性。
    """
    try:
        result = await ms_agent_service.run_deep_research(
            query=request.query,
            model=request.model,
            exa_api_key=request.exa_api_key,
            serpapi_api_key=request.serpapi_api_key,
            max_rounds=request.max_rounds,
        )
        return DeepResearchResponse(**result)
    except Exception as e:
        logger.error(f"start_deep_research error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get(
    "/research/{task_id}",
    response_model=MSAgentTaskStatus,
    summary="[ms-agent] 查询深度研究任务状态 + 读取报告",
    tags=["msagent"],
)
async def get_research_status(task_id: str):
    """
    查询深度研究任务的当前状态，并在完成后返回 `report` 字段（full Markdown）。

    `status` 枚举：
    - `pending`    — 目录已建，进程尚未产出文件
    - `running`    — 进程运行中
    - `completed`  — 正常结束（returncode=0）
    - `failed`     — 异常退出（见 `stderr_tail` 了解原因）
    - `not_found`  — 任务 ID 不存在
    """
    status = ms_agent_service.get_task_status("research", task_id)
    return MSAgentTaskStatus(**status)


@router.get(
    "/research",
    summary="[ms-agent] 列出所有深度研究任务",
    tags=["msagent"],
)
async def list_research_tasks():
    """返回所有深度研究任务的摘要列表（按创建时间倒序）。"""
    tasks = ms_agent_service.list_tasks("research")
    return {"tasks": tasks, "total": len(tasks)}


@router.post(
    "/code/generate",
    response_model=CodeGenResponse,
    summary="[ms-agent] 触发 code_genesis 代码生成任务",
    tags=["msagent"],
)
async def start_code_genesis(request: CodeGenRequest):
    """
    启动一个 **code_genesis** 复杂代码生成任务（异步 fire-and-forget）。

    工作流（7 阶段 DAG）：
    > user_story → architect → file_design → file_order → install → coding → refine

    输出（容器内 `/app/data/outputs/code/{task_id}/output/`）：
    - 完整可运行项目目录（含依赖安装脚本、README.md 等）

    立即返回 `task_id`，通过 `GET /code/{task_id}` 轮询状态。

    适用场景：
    - 多文件、多模块的复杂系统开发
    - 需要自动设计文件结构的项目
    - 需要依赖安装和代码精炼的完整工作流

    > 注意：任务运行时间通常 10-60 分钟，取决于项目复杂度。
    """
    try:
        result = await ms_agent_service.run_code_genesis(
            query=request.query,
            model=request.model,
        )
        return CodeGenResponse(**result)
    except Exception as e:
        logger.error(f"start_code_genesis error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get(
    "/code/{task_id}",
    response_model=MSAgentTaskStatus,
    summary="[ms-agent] 查询代码生成任务状态 + 产物列表",
    tags=["msagent"],
)
async def get_code_status(task_id: str):
    """
    查询代码生成任务的当前状态，在完成后通过 `output_files` 返回所有产物文件路径。
    """
    status = ms_agent_service.get_task_status("code", task_id)
    return MSAgentTaskStatus(**status)


@router.get(
    "/code",
    summary="[ms-agent] 列出所有代码生成任务",
    tags=["msagent"],
)
async def list_code_tasks():
    """返回所有代码生成任务的摘要列表（按创建时间倒序）。"""
    tasks = ms_agent_service.list_tasks("code")
    return {"tasks": tasks, "total": len(tasks)}


# ============================================================
# Phase 11 — 代码语义索引 /code-index/*
# ============================================================

from app.models.schemas import (
    CodeSearchRequest,
    CodeSearchResponse,
    CodeSymbolResponse,
    IndexStatusResponse,
)


@router.post(
    "/code-index/search",
    response_model=CodeSearchResponse,
    summary="语义搜索代码符号",
    tags=["code-index"],
)
async def search_code(request: CodeSearchRequest):
    """
    通过自然语言查询定位项目代码中的函数/类/方法。

    返回文件路径 + 精确行号范围 + 函数签名，
    可直接用于 read_file(file_path, line_start, line_end) 精准读取代码片段，
    大幅降低 AI 上下文 Token 消耗。

    示例请求：
    ```json
    {"query": "处理 PDF 上传并校验文件大小", "top_k": 3}
    ```
    """
    results = await code_indexer.search(
        query=request.query,
        top_k=request.top_k,
        symbol_types=request.symbol_types,
        file_pattern=request.file_pattern,
    )
    status = code_indexer.get_index_status()
    return CodeSearchResponse(
        results=[CodeSymbolResponse(**r) for r in results],
        total_indexed=status["total_symbols"],
        query=request.query,
    )


@router.post(
    "/code-index/rebuild",
    summary="全量重建代码索引",
    tags=["code-index"],
)
async def rebuild_code_index():
    """
    清空并重建整个代码语义索引（全量模式）。

    全流程 AST 解析，零 LLM Token 消耗，通常 200-500ms 内完成。
    建议：修改大量文件或首次部署后调用；日常修改由增量更新覆盖。
    """
    stats = await code_indexer.build_full_index()
    return {"status": "ok", "stats": stats}


@router.post(
    "/code-index/update",
    summary="增量更新代码索引",
    tags=["code-index"],
)
async def update_code_index():
    """
    检测文件变更并增量更新索引（仅处理 hash 变化的文件）。

    零 LLM Token 消耗，通常 < 20ms。
    """
    stats = await code_indexer.update_incremental()
    return {"status": "ok", "stats": stats}


@router.get(
    "/code-index/status",
    response_model=IndexStatusResponse,
    summary="查询代码索引状态",
    tags=["code-index"],
)
async def get_code_index_status():
    """返回索引统计信息：文件数、符号数、过期文件数、最后索引时间。"""
    status = code_indexer.get_index_status()
    return IndexStatusResponse(**status)


@router.get(
    "/code-index/symbols",
    summary="列出指定文件的所有符号",
    tags=["code-index"],
)
async def list_file_symbols(file_path: str):
    """
    列出指定文件中的所有已索引符号（按行号排序）。

    用于 AI 快速了解一个文件的结构概览（类似「目录页」），
    无需读取整个文件内容，极大节省 Token。

    示例：GET /code-index/symbols?file_path=app/services/worker.py
    """
    with Session(engine) as session:
        symbols = session.exec(
            select(CodeSymbol)
            .where(CodeSymbol.file_path == file_path)
            .order_by(CodeSymbol.line_start)
        ).all()
    return {
        "file_path": file_path,
        "symbol_count": len(symbols),
        "symbols": [
            {
                "symbol_name": s.symbol_name,
                "symbol_type": s.symbol_type,
                "parent_symbol": s.parent_symbol,
                "line_start": s.line_start,
                "line_end": s.line_end,
                "signature": s.signature,
            }
            for s in symbols
        ],
    }
